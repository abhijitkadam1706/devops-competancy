#!/usr/bin/env python3
"""
Generate PRCICD-001 & PRCICD-002 Compliance Document (DOCX)
mern-auth CI/CD — Jenkins Master/Slave | K8s | Prometheus
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Colour helpers ──────────────────────────────────────────────────────────────
def hex2rgb(h): r=int(h[1:3],16); g=int(h[3:5],16); b=int(h[5:7],16); return RGBColor(r,g,b)

NAVY   = hex2rgb('#1E3A5F')
GREEN  = hex2rgb('#064E3B')
PURPLE = hex2rgb('#1E1B4B')
SLATE  = hex2rgb('#1E293B')
WHITE  = RGBColor(0xFF,0xFF,0xFF)
GOLD   = hex2rgb('#B45309')
LGRAY  = hex2rgb('#F8FAFC')
BORDER = hex2rgb('#E2E8F0')

def set_cell_bg(cell, hex_color):
    tc  = cell._tc
    tcp = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color.lstrip('#'))
    tcp.append(shd)

def cell_text(cell, text, bold=False, size=10, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    para = cell.paragraphs[0]
    para.alignment = align
    run  = para.add_run(text)
    run.bold      = bold
    run.font.size = Pt(size)
    if color: run.font.color.rgb = color

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13 if level == 1 else 11)
    run.font.color.rgb = NAVY if level == 1 else GREEN
    p.paragraph_format.space_before = Pt(14 if level==1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_para(doc, text, size=10, color=None, space_after=6, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size    = Pt(size)
    run.font.italic  = italic
    if color: run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p

# ── Document setup ──────────────────────────────────────────────────────────────
doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# Default font
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10)

# ── TITLE PAGE ──────────────────────────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_p.add_run('End-to-End CI/CD Pipeline')
tr.bold           = True
tr.font.size      = Pt(22)
tr.font.color.rgb = NAVY

doc.add_paragraph()
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub_p.add_run('mern-auth Application — DevOps Competency Submission')
sr.font.size      = Pt(13)
sr.font.color.rgb = hex2rgb('#475569')

doc.add_paragraph()
badge_p = doc.add_paragraph()
badge_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
b1 = badge_p.add_run('  PRCICD-001: Software Release Workflows  ')
b1.bold           = True
b1.font.size      = Pt(10)
b1.font.color.rgb = WHITE
b1.font.highlight_color = None

b2 = badge_p.add_run('   ')
b3 = badge_p.add_run('  PRCICD-002: Build and Test Code  ')
b3.bold           = True
b3.font.size      = Pt(10)
b3.font.color.rgb = WHITE

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta_p.add_run(f'Prepared: {datetime.date.today().strftime("%d %B %Y")}  ·  Architecture: Jenkins Master/Slave  ·  Runtime: Kubernetes  ·  Monitoring: Prometheus')
mr.font.size      = Pt(8.5)
mr.font.italic    = True
mr.font.color.rgb = hex2rgb('#94A3B8')

doc.add_paragraph()

# ── SECTION 1 — OVERVIEW ────────────────────────────────────────────────────────
add_heading(doc, '1. Overview')
add_para(doc,
    'This document describes the end-to-end CI/CD pipeline designed for the mern-auth '
    'application (MERN stack). It maps each pipeline stage to the two DevOps competency '
    'controls under review, demonstrating how the implementation satisfies both release '
    'governance (PRCICD-001) and build quality (PRCICD-002) requirements.')

add_para(doc,
    'The pipeline runs on a Jenkins Master/Slave architecture deployed within a Kubernetes '
    'cluster. The Master coordinates the workflow while ephemeral slave agents handle compute-'
    'intensive tasks. Prometheus provides real-time production monitoring and triggers automatic '
    'rollback when error thresholds are exceeded.', italic=False)

# ── SECTION 2 — ARCHITECTURE DIAGRAM ────────────────────────────────────────────
add_heading(doc, '2. Architecture Diagram')
add_para(doc, 'The diagram below shows the three-lane pipeline structure: Jenkins Master (orchestrator), Jenkins Slaves (workers), and External Systems.',
         size=9, color=hex2rgb('#64748B'))

try:
    doc.add_picture(
        r'c:\Users\Abhijit kadam\Desktop\Cloude\DevOps Competency-2\cicd_architecture_final.png',
        width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
except Exception as e:
    add_para(doc, f'[Architecture diagram: cicd_architecture_final.png — {e}]',
             italic=True, color=hex2rgb('#EF4444'))

# ── SECTION 3 — MASTER / SLAVE OVERVIEW ─────────────────────────────────────────
add_heading(doc, '3. Master / Slave Design')
ms_data = [
    ('Jenkins Master', 'Permanent EC2/K8s pod in the cicd namespace',
     'Receives webhook on PR merge, schedules agent pods, holds the Manual Approval Gate, '
     'sends Slack notifications, stores the audit trail. Never runs build or test work itself.'),
    ('Build Agent (Slave)', 'Ephemeral K8s pod — label: build-agent',
     'Runs compile, unit tests, SonarQube SAST, Docker image build, Trivy SCA scan, and '
     'registry push. Terminated after each job to eliminate state bleed between runs.'),
    ('Test Agent (Slave)', 'Ephemeral K8s pod — label: test-agent',
     'Runs integration tests via Newman (Postman collections) and OWASP ZAP DAST against '
     'a live container. Also queries Prometheus for the post-production health check.'),
    ('Deploy Agent (Slave)', 'Ephemeral K8s pod — label: deploy-agent',
     'Executes kubectl commands to update Kubernetes Deployments in staging and production '
     'namespaces. Handles Blue/Green traffic switching by patching the Service selector.'),
]
tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0].cells
hdrs = ['Component', 'Location', 'Responsibility']
for i, h in enumerate(hdrs):
    set_cell_bg(hdr[i], '#1E3A5F')
    cell_text(hdr[i], h, bold=True, color=WHITE, size=9)
for row_d in ms_data:
    row = tbl.add_row().cells
    for i, txt in enumerate(row_d):
        cell_text(row[i], txt, size=9)
        if i % 2 == 0:
            set_cell_bg(row[i], '#F1F5F9')
doc.add_paragraph()

# ── SECTION 4 — STAGE-BY-STAGE BREAKDOWN ────────────────────────────────────────
add_heading(doc, '4. Pipeline Stages')
add_para(doc, 'Each stage is mapped to the control it satisfies, the tools used, and the evidence artefact produced.', size=9, italic=True)

stages = [
    # (stage, slave, control, tools, what_happens, evidence)
    ('1\nCheckout & Compile', 'Build Agent', 'PRCICD-002',
     'Git · npm ci · React build',
     'Pipeline triggers on PR merge to main. Build agent clones the repo and compiles '
     'both backend (Node.js) and frontend (React), confirming the codebase builds cleanly.',
     'Build log with exit code 0 and npm install manifest'),

    ('2–3\nUnit Tests & SAST', 'Build Agent', 'PRCICD-002',
     'Mocha · Jest · SonarQube',
     'Unit tests run across backend and frontend suites; JUnit XML reports are published. '
     'SonarQube scanner analyses code for injection flaws, insecure patterns, and anti-patterns. '
     'A quality gate check enforces ≥ 80% line coverage and blocks the pipeline on failure.',
     'JUnit test report · SonarQube dashboard · coverage %'),

    ('4\nDocker Build & SCA', 'Build Agent', 'PRCICD-002',
     'Docker 24 · Trivy 0.51',
     'Source is containerised. Trivy scans the image for known CVEs in OS packages and '
     'application dependencies. A SBOM in CycloneDX format is generated. Zero critical CVEs '
     'are permitted; any finding halts the pipeline immediately.',
     'trivy-report.json · sbom.json · Docker image digest'),

    ('5\nPush Stage Registry', 'Build Agent', 'PRCICD-002',
     'Docker registry',
     'Verified image is pushed to the stage registry tagged with the short commit SHA '
     '(immutable) and :latest. This is the exact binary that will travel through every '
     'downstream environment — no rebuilds between stages.',
     'Registry push log · image tag = commit SHA'),

    ('6–7\nIntegration Tests & DAST', 'Test Agent', 'PRCICD-002',
     'Newman · OWASP ZAP',
     'A live container is started on the test agent. Newman executes Postman API collections '
     'verifying auth endpoints, token flows, and error codes. OWASP ZAP then runs a baseline '
     'DAST scan, simulating real-world attacks against the running application. Zero high-'
     'severity findings are required to proceed.',
     'integration-results.xml · zap-report.html · zap-report.xml'),

    ('8\nDeploy to Staging (K8s)', 'Deploy Agent', 'PRCICD-001',
     'kubectl · Kubernetes',
     'Deploy agent runs kubectl set image against the mern-auth-staging namespace. '
     'Kubernetes performs a rolling update; rollout status is polled until complete. '
     'A curl smoke test confirms the /health endpoint returns HTTP 200 before the gate opens.',
     'kubectl rollout output · smoke test log'),

    ('9\nManual Approval Gate', 'Master (no agent)', 'PRCICD-001',
     'Jenkins Input Step · Slack',
     'Pipeline pauses. An SNS/Slack notification is sent to the release manager with links '
     'to the SonarQube report, Trivy CVE report, ZAP DAST report, and test coverage. '
     'The approver reviews all evidence and clicks Approve or Reject. The approver identity '
     'and timestamp are captured in the Jenkins build record.',
     'Approval record: approver name + timestamp in build history'),

    ('10\nPush Prod Registry', 'Build Agent', 'PRCICD-001',
     'Docker registry',
     'Staging-validated image is retagged and pushed to the separate prod registry '
     '(registry.internal/mern-auth-prod) with the commit SHA tag and :stable. '
     'Separate registries provide an audit boundary between non-production and production images.',
     'Prod registry push log · :stable tag'),

    ('11\nProd Deploy — Blue/Green', 'Deploy Agent', 'PRCICD-001',
     'kubectl · Blue/Green strategy',
     'The new image is deployed to the Green Deployment in the mern-auth-production namespace. '
     'Once all Green pods pass readiness probes, the mern-auth-prod Service selector is patched '
     'to slot=green, shifting 100% of traffic with zero downtime. Blue remains live and idle '
     'for instant rollback.',
     'kubectl patch output · Service selector change log'),

    ('12\nPrometheus Monitor & Rollback', 'Test Agent', 'PRCICD-001',
     'Prometheus API',
     'For five minutes after traffic switch, the test agent polls Prometheus every 30 seconds '
     'for the production HTTP 5xx error rate. If the rate exceeds 5%, the pipeline automatically '
     'patches the Service selector back to slot=blue, reverting to the previous stable version '
     'with no human intervention required.',
     'Prometheus query results · rollback log (if triggered)'),
]

for s in stages:
    stage_no, slave, ctrl, tools, what, evidence = s
    tbl2 = doc.add_table(rows=5, cols=2)
    tbl2.style = 'Table Grid'
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row (stage name spanning 2 cols)
    hrow = tbl2.rows[0]
    hrow.cells[0].merge(hrow.cells[1])
    bg = '#1E3A5F' if ctrl == 'PRCICD-001' else '#064E3B'
    set_cell_bg(hrow.cells[0], bg)
    cell_text(hrow.cells[0], f'Stage {stage_no}', bold=True, color=WHITE, size=9.5)

    # Detail rows
    rows_data = [
        ('Agent', slave),
        ('Control', ctrl),
        ('Tools', tools),
        ('What happens', what),
        ('Evidence', evidence),
    ]
    # Merge header then fill 4 data rows
    data_rows = [
        ('Agent / Control', f'{slave}  ·  {ctrl}'),
        ('Tools', tools),
        ('What happens', what),
        ('Evidence artefact', evidence),
    ]
    for i, (lbl, val) in enumerate(data_rows):
        r = tbl2.rows[i + 1]
        set_cell_bg(r.cells[0], '#F1F5F9')
        cell_text(r.cells[0], lbl, bold=True, size=8.5)
        cell_text(r.cells[1], val, size=8.5)

    doc.add_paragraph()

# ── SECTION 5 — EVIDENCE SUMMARY ────────────────────────────────────────────────
add_heading(doc, '5. Evidence Summary')
add_para(doc, 'The following artefacts should be attached to the APN competency case study submission:', size=9)

ev_tbl = doc.add_table(rows=1, cols=3)
ev_tbl.style = 'Table Grid'
ev_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Control', 'Artefact', 'Source']):
    set_cell_bg(ev_tbl.rows[0].cells[i], '#1E293B')
    cell_text(ev_tbl.rows[0].cells[i], h, bold=True, color=WHITE, size=9)

ev_rows = [
    ('PRCICD-001', 'Multibranch Pipeline definition + branching strategy',       'Jenkins → Configure → Pipeline script'),
    ('PRCICD-001', 'Manual approval record (approver + timestamp)',               'Jenkins Build #N → Input step log'),
    ('PRCICD-001', 'Blue/Green deployment log + traffic switch',                 'kubectl output archived in Jenkins'),
    ('PRCICD-001', 'Prometheus rollback log (or "no rollback" confirmation)',     'Pipeline Stage 12 console log'),
    ('PRCICD-002', 'buildspec / Jenkinsfile showing compile + test commands',    'Jenkinsfile in repo root'),
    ('PRCICD-002', 'JUnit test report (unit + integration)',                     'Jenkins Test Results page'),
    ('PRCICD-002', 'SonarQube quality gate result (≥ 80% coverage)',             'SonarQube dashboard export'),
    ('PRCICD-002', 'Trivy CVE scan report (trivy-report.json)',                  'Jenkins Artifacts tab'),
    ('PRCICD-002', 'SBOM — CycloneDX format (sbom.json)',                        'Jenkins Artifacts tab'),
    ('PRCICD-002', 'OWASP ZAP DAST report (zap-report.html)',                    'Jenkins Artifacts tab'),
    ('PRCICD-002', 'ECR/Registry image push log with commit SHA tag',            'Jenkins Stage 5 / Stage 10 log'),
]
for ctrl, art, src in ev_rows:
    r = ev_tbl.add_row().cells
    bg = '#DBEAFE' if ctrl == 'PRCICD-001' else '#D1FAE5'
    set_cell_bg(r[0], bg)
    cell_text(r[0], ctrl, bold=True, size=8.5)
    cell_text(r[1], art,  size=8.5)
    cell_text(r[2], src,  size=8.5, color=hex2rgb('#475569'))

doc.add_paragraph()

# ── SECTION 6 — TOOL RATIONALE ───────────────────────────────────────────────────
add_heading(doc, '6. Tool Selection Rationale')
tools_data = [
    ('Jenkins Master/Slave', 'Industry-standard orchestrator with Kubernetes cloud plugin. '
     'Master/Slave separation ensures the controller is never a build bottleneck; ephemeral '
     'slaves eliminate state contamination between pipeline runs.'),
    ('Kubernetes (K8s)', 'Provides declarative, version-controlled infrastructure for agent '
     'pods, staging, and production environments. Blue/Green deployments are implemented natively '
     'via Deployment slot labels and a single Service selector patch.'),
    ('SonarQube', 'SAST tool with built-in quality gate API consumed by the Jenkins '
     '"waitForQualityGate" step. Provides coverage thresholds and vulnerability rules that '
     'directly meet the PRCICD-002 code quality criterion.'),
    ('Trivy', 'Fast, lightweight SCA/CVE scanner with SBOM generation (CycloneDX). '
     'Supports image scanning, exit-code enforcement, and JSON output for pipeline integration.'),
    ('OWASP ZAP', 'DAST scanner that tests the running application for OWASP Top 10 vulnerabilities, '
     'complementing Trivy\'s static analysis. Baseline scan mode is CI-friendly with a defined exit code.'),
    ('Newman', 'CLI runner for Postman collections. Provides repeatable, version-controlled '
     'API integration tests with JUnit-format output natively supported by Jenkins.'),
    ('Prometheus', 'Pull-based metrics collection from both the Jenkins cluster and application '
     'pods. Alert rules enforce a 5% error rate SLA; the test agent queries the Prometheus API '
     'to implement automated rollback logic without external dependencies.'),
]
tool_tbl = doc.add_table(rows=1, cols=2)
tool_tbl.style = 'Table Grid'
tool_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Tool', 'Why Used']):
    set_cell_bg(tool_tbl.rows[0].cells[i], '#1E293B')
    cell_text(tool_tbl.rows[0].cells[i], h, bold=True, color=WHITE, size=9)
for tname, treason in tools_data:
    r = tool_tbl.add_row().cells
    set_cell_bg(r[0], '#F8FAFC')
    cell_text(r[0], tname, bold=True, size=8.5)
    cell_text(r[1], treason, size=8.5)

doc.add_paragraph()

# ── FOOTER ───────────────────────────────────────────────────────────────────────
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer_p.add_run('AWS DevOps Competency Submission  ·  mern-auth CI/CD  ·  PRCICD-001 / PRCICD-002')
fr.font.size = Pt(8)
fr.font.italic = True
fr.font.color.rgb = hex2rgb('#94A3B8')

# ── Save ─────────────────────────────────────────────────────────────────────────
out = r'c:\Users\Abhijit kadam\Desktop\Cloude\DevOps Competency-2\CICD_PRCICD_Compliance_Document.docx'
doc.save(out)
print(f"Document saved → {out}")
